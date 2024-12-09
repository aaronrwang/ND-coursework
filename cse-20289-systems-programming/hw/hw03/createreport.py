#Ethan Little elittle2@nd.edu
#Derick Shi dhsi2@nd.edu
#Aaron Wang awang27@nd.edu

import argparse
from docx import Document
from docx.shared import Inches

def read_text_file(file_path):
    """Read content from a text file."""
    with open(file_path, 'r') as file:
        return file.read()

def create_word_document(textFile, image_path_list, statistics, output_filename):
    """Create a Word document with the provided text, statistics table, and image."""
    # Initialize a new Document
    doc = Document()
    text = read_text_file(textFile)
    # Add the text content
    doc.add_paragraph(text)

    # Create and add the statistics table
    add_statistics_table(doc, statistics)

    # Add the image
    
    for graph in image_path_list: #make graph for every month as needed
        doc.add_paragraph(graph[:7])
        doc.add_picture(graph, width=Inches(5))  # Adjust width as needed

    # Save the document
    doc.save(output_filename)

def add_statistics_table(doc, statistics):
    """Create a table in the document using the statistics dictionary."""
    # Create a table with 1 header row and the appropriate number of data rows
    table = doc.add_table(rows=1, cols=2)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Statistic'
    hdr_cells[1].text = 'Value'

    # Populate the table with statistics
    for key, value in statistics.items():
        row_cells = table.add_row().cells
        row_cells[0].text = key
        row_cells[1].text = str(value)

if __name__ == "__main__":
    parser = argparse.ArgumentParser(description='Generate a Word document from a text file and an image.')
    parser.add_argument('text_file', type=str, help='Path to the input text file.')
    parser.add_argument('image_file', type=str, help='Path to the input PNG image file.')
    parser.add_argument('output_file', type=str, help='Path to the output Word document.')

    args = parser.parse_args()

    # Read text from the text file
    text_content = read_text_file(args.text_file)

    # Dummy statistics with all entries as zeroes
    statistics = {
        'Statistic A': 0,
        'Statistic B': 0,
        'Statistic C': 0,
        'Statistic D': 0,
    }

    # Create the Word document
    create_word_document(text_content, args.image_file, statistics, args.output_file)